import os
from docx import Document
from django.conf import settings


def _find_borrador_docx() -> str:
    """
    Devuelve la ruta absoluta a templates/base/borrador.docx buscando en:
    - settings.TEMPLATES[*]['DIRS']
    - <BASE_DIR>/templates
    """
    candidates = []

    # 1) DIRS configurados en settings.TEMPLATES
    for cfg in getattr(settings, "TEMPLATES", []):
        for d in cfg.get("DIRS", []):
            candidates.append(os.path.join(d, "base", "borrador.docx"))

    # 2) Fallback: <BASE_DIR>/templates/base/borrador.docx
    candidates.append(os.path.join(settings.BASE_DIR, "templates", "base", "borrador.docx"))

    for path in candidates:
        if os.path.exists(path):
            return path

    raise FileNotFoundError(
        "No se encontró 'borrador.docx'. Se intentó en:\n" + "\n".join(candidates)
    )


def fill_word_template(context: dict, output_path: str) -> str:
    """
    Llena la plantilla 'borrador.docx' con los valores de `context` usando marcadores {{clave}}.
    Guarda el resultado en `output_path` y devuelve esa misma ruta.
    """
    template_path = _find_borrador_docx()
    doc = Document(template_path)

    def _replace_text_in_text(text: str) -> str:
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"  # -> {{clave}}
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        return text

    # Reemplazo en párrafos
    for p in doc.paragraphs:
        # Nota: usar p.text sobrescribe estilos del párrafo si los placeholders están partidos en runs.
        # Mantener simple por ahora (plantilla con placeholders en un mismo run).
        p.text = _replace_text_in_text(p.text)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = _replace_text_in_text(cell.text)

    # Guardar en la ruta de salida
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
